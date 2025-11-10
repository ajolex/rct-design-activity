from __future__ import annotations

import html
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PREFERRED_TABLE_STYLE = "Grid Table 4 Accent 2"
FALLBACK_TABLE_STYLE = "Normal Table"

PARTICIPANT_GUIDANCE = [
    "Move in order but keep the pace brisk. Aim to make decisions, not perfect prose.",
    "Write directly in the boxes. Add sticky notes or sketches where helpful.",
    "Return to prompts marked with [ ] if you need to revisit them during the gallery walk.",
]

PARTICIPANT_STEPS = [
    {
        "number": 1,
        "title": "Frame the Challenge",
        "goal": "Clarify the problem you are tackling and the change you want to see.",
        "actions": [
            "Review the program card and underline the core challenge.",
            "Note the primary participants and delivery setting.",
            "Write a one-sentence statement of success for the next 12 months.",
        ],
        "tip": "Keep your problem statement specific to one population and one outcome.",
        "notes": [
            {"label": "Program title", "lines": 1},
            {"label": "Target group", "lines": 1},
            {"label": "Delivery setting", "lines": 1},
            {"label": "Success in 12 months looks like", "lines": 3},
        ],
    },
    {
        "number": 2,
        "title": "Map the Theory of Change",
        "goal": "Connect activities to outcomes so your randomization follows the logic.",
        "actions": [
            "List the major activities and the immediate outputs you can measure.",
            "Highlight the outcomes that must shift before your long-term impact appears.",
            "Flag the assumptions that feel riskiest or hardest to prove.",
        ],
        "tip": "If you cannot draw a tight line from activity to outcome, consider narrowing scope.",
        "notes": [
            {"label": "Riskiest assumption", "lines": 3},
            {"label": "Early signal to watch", "lines": 3},
        ],
    },
    {
        "number": 3,
        "title": "Design Measurement",
        "goal": "Choose indicators and instruments that capture change credibly.",
        "actions": [
            "Lock one primary outcome and the metric you will use.",
            "Select instruments that are realistic for your team to field.",
            "Set timing for baseline and follow-up so you capture change.",
        ],
        "tip": "If an indicator feels fuzzy, add a quick definition or example in the notes.",
        "notes": [
            {"label": "Primary outcome definition", "lines": 3},
            {"label": "Instrument(s)", "lines": 3},
        ],
    },
    {
        "number": 4,
        "title": "Plan Randomization",
        "goal": "Select an assignment approach and sanity check for implementation.",
        "actions": [
            "Pick the level where you will randomize (individuals, villages, classrooms, clinics, etc.).",
            "List the steps to assign units and keep the process transparent.",
            "Anticipate spillovers or compliance risks and note how you will manage them.",
        ],
        "tip": "If spillovers feel unavoidable, consider cluster-level assignment.",
        "notes": [
            {"label": "Randomization unit", "lines": 2},
            {"label": "Method", "lines": 2},
            {"label": "Checks we will run", "lines": 4},
        ],
    },
    {
        "number": 5,
        "title": "Safeguard Implementation",
        "goal": "Make the plan operational with rhythms, resources, and quick signals.",
        "actions": [
            "Establish communication protocols to synchronize program rollout, monitoring, and data collection",
            "Capture logistics that could slow you down and draft mitigation steps.",
            "Select an interim indicator or dashboard you will review during rollout.",
        ],
        "tip": "Assign a lead person to each risk so follow-up happens quickly.",
        "notes": [
            {"label": "Team check-ins", "lines": 3},
            {"label": "Risks to watch", "lines": 3},
        ],
    },
    {
        "number": 6,
        "title": "Decide and Commit",
        "goal": "Record the trigger that will guide your next decision once data arrives.",
        "actions": [
            "State the metric and threshold that signals scale, revise, or stop.",
            "Note who needs to see the results and how they prefer updates.",
        ],
        "tip": "Be concrete: threshold, time frame, and decision owner.",
        "notes": [
            {"label": "Decision trigger", "lines": 3},
            {"label": "Stakeholders to brief", "lines": 3},
            {"label": "Next steps after briefing", "lines": 3},
        ],
    },
]

TOC_GRID_CELLS = [
    {"title": "Needs", "prompt": "What is the core problem/need?", "lines": 4},
    {"title": "Inputs", "prompt": "Resources and/or partners you bring to the program.", "lines": 4},
    {"title": "Activities", "prompt": "Key actions or services you will deliver.", "lines": 4},
    {"title": "Outputs", "prompt": "Products or services produced by program activities; deliverables.", "lines": 4},
    {"title": "Intermediate Outcomes", "prompt": "Mechanisms linking early changes to impact.", "lines": 4},
    {"title": "Long-Term Impact", "prompt": "Ultimate change you are aiming to achieve.", "lines": 4},
]

TOC_NOTE_CELLS = [
    {"title": "Critical Assumptions", "prompt": "What must stay true? List the riskiest points.", "lines": 3},
    {"title": "Rapid Validation Ideas", "prompt": "Quick tests to check those assumptions.", "lines": 3},
]

MEASUREMENT_HEADERS = [
    "Construct",
    "Indicator and Metric",
    "Instrument or Data Source",
    "Timing and Frequency",
    "Notes",
]

MEASUREMENT_CONSTRUCTS = [
    {"label": "Primary Outcome", "rows": 4},
    {"label": "Key Mechanism", "rows": 3},
    {"label": "Implementation Health", "rows": 3},
]

RANDOMIZATION_HEADERS = [
    "Randomization Level",
    "Method",
    "Assignment Steps and Checks",
    "Spillover Mitigation",
    "Power Considerations",
]

CHECKLIST_BLOCKS = [
    ("Data Quality", "Quick validations, back checks, or monitoring ideas."),
    ("Risks and Mitigation", "Compliance risks, attrition threats, mitigation tactics."),
    ("Decision Trigger", "Metric and threshold that signals scale or redesign."),
]

DATA_QUALITY_ROWS = [
    {
        "title": "Team Rhythm",
        "items": ["Daily huddle plan", "Data entry cadence", "Escalation contact"],
    },
    {
        "title": "Field Logistics",
        "items": ["Materials packed", "Travel or connectivity plan", "Budget guardrails"],
    },
    {
        "title": "Signal Watch",
        "items": ["Interim indicator to review", "Decision trigger lead", "Follow-up plan"],
    },
]

REFLECTION_PROMPTS = [
    {"label": "Biggest assumption we still need to verify", "lines": 3},
    {"label": "Evidence that will convince a decision maker", "lines": 3},
    {"label": "Support we need from others", "lines": 3},
]

SPRINT_CHECKLIST = [
    "Primary outcome and indicator locked in",
    "Randomization unit clear and feasible",
    "Measurement instruments matched to outcomes",
    "Assumptions flagged with validation ideas",
    "Decision trigger recorded and owner assigned",
]

SESSION_SNAPSHOT_ROWS = [
    ("Welcome Spark", "Why randomized trials matter right now", "Facilitator briefing", "4 min"),
    ("Design Sprint", "Teams work through the workbook", "Team breakout", "18 min"),
    ("Gallery Feedback", "Peer review and refinement", "Walk-through", "5 min"),
    ("Commit to Next Step", "Teams capture decisions and owners", "Team huddle", "3 min"),
]

FACILITATOR_PREP = [
    "Print one participant workbook per person in landscape, double sided if possible.",
    "Prepare program cards, measurement planner, and theory of change posters.",
    "Set up a visible timer, sticky notes, and markers at each table.",
    "Post gallery prompts: Alignment, Feasibility, Decision Trigger.",
]

FACILITATOR_FLOW = [
    (
        "Welcome Spark (4 minutes)",
        [
            "Share a short story where randomization clarified a program decision.",
            "Frame the sprint goal: produce a defendable RCT plan using the workbook.",
            "Preview the three tangible outputs: theory of change, measurement plan, decision trigger.",
        ],
    ),
    (
        "Design Sprint (18 minutes)",
        [
            "Prompt teams to divide roles: evidence lead, logician, skeptic, recorder.",
            "Check in at minutes 5, 10, and 15 with catalytic questions.",
            "Encourage teams to mark blanks they will revisit rather than stalling.",
        ],
    ),
    (
        "Gallery Feedback (5 minutes)",
        [
            "Teams post key pages on the wall for a gallery walk.",
            "Neighbors leave one sticky per prompt: Alignment, Feasibility, Decision Trigger.",
        ],
    ),
    (
        "Commit to Next Step (3 minutes)",
        [
            "Teams capture final tweaks and nominate a spokesperson.",
            "Invite one-sentence commitments during the share-out.",
        ],
    ),
]

COACHING_PROMPTS = [
    "Where could spillovers or compliance issues weaken this randomization?",
    "Which assumption, if wrong, breaks the theory of change?",
    "How will you measure the primary outcome without overloading the team?",
    "What decision will you make once the first results arrive?",
]

DEBRIEF_BULLETS = [
    "Collect workbooks or photograph key pages for documentation.",
    "Invite teams to share their decision trigger in one sentence.",
    "Offer optional clinics on power, budgeting, or ethics for interested teams.",
    "Share digital copies of all materials within 24 hours.",
]

QUICK_REFERENCE_CUES = [
    ("Time slipping", "Announce checkpoints and nudge teams to advance."),
    ("Teams stuck", "Ask for their riskiest assumption and how they will test it."),
    ("Over-scoping", "Remind them to lock a single primary outcome and method."),
    ("Wrapping up", "Capture decision triggers and reset the room."),
]

PARTICIPANT_HTML_CSS = """
:root{--primary:#164a7f;--accent:#2fa6dc;--ink:#1f2d3d;--surface:#ffffff;}
*{box-sizing:border-box;}
body{margin:0;font-family:'Segoe UI','Helvetica Neue',Arial,sans-serif;background:linear-gradient(135deg,#f4f8ff,#ffffff);color:var(--ink);} 
main{margin:0;padding:0;background:var(--surface);} 
.cover-page{width:100vw;height:100vh;display:flex;align-items:center;justify-content:center;background:var(--surface);page-break-after:always;} 
.cover-page img{width:100%;height:100%;object-fit:cover;} 
.content{max-width:1040px;margin:0 auto;padding:48px 56px;page-break-after:always;} 
h1{margin:0 0 12px;color:var(--primary);font-size:2.35rem;letter-spacing:0.05em;} 
h2{margin:36px 0 14px;color:var(--primary);font-size:1rem;letter-spacing:0.08em;text-transform:uppercase;} 
p{line-height:1.65;margin:0 0 1.1em;} 
ul.guidance{margin:0 0 28px;padding-left:1.5rem;} 
.step-page{page-break-before:always;min-height:100vh;padding:48px 56px;background:var(--surface);} 
.step-wrapper{max-width:980px;margin:0 auto;height:100%;display:flex;flex-direction:column;} 
.step-header{display:flex;flex-direction:column;gap:12px;margin-bottom:28px;} 
.step-title{margin:0;font-size:2rem;color:var(--primary);letter-spacing:0.06em;text-transform:uppercase;} 
.step-goal{font-size:1.05rem;color:rgba(31,45,61,0.78);} 
.step-body{display:flex;flex-direction:row;gap:32px;flex:1;} 
.step-actions{flex:1.08;background:rgba(22,74,127,0.06);border-radius:18px;padding:24px 26px;box-shadow:inset 0 0 0 1px rgba(22,74,127,0.12);} 
.step-actions ul{margin:0;padding-left:1.25rem;font-size:1.05rem;} 
.step-actions li{margin-bottom:0.7rem;} 
.tip{margin-top:18px;padding:12px 16px;border-left:4px solid var(--accent);background:rgba(47,166,220,0.12);border-radius:14px;font-size:0.95rem;} 
.note-stack{flex:1.05;display:grid;gap:18px;} 
.note-card{background:#fff;border-radius:18px;padding:18px 20px;border:1px solid rgba(22,74,127,0.18);box-shadow:0 24px 42px rgba(22,74,127,0.16);display:flex;flex-direction:column;gap:10px;} 
.note-label{font-size:0.82rem;letter-spacing:0.08em;text-transform:uppercase;color:rgba(31,45,61,0.7);} 
.note-input,.note-area{width:100%;border:1px dashed rgba(22,74,127,0.4);border-radius:12px;padding:12px;font-family:inherit;font-size:1rem;background:rgba(244,248,255,0.65);} 
.note-input{height:52px;} 
.note-area{min-height:120px;resize:vertical;} 
.step-extension{margin-top:34px;background:rgba(22,74,127,0.06);border-radius:22px;padding:28px 30px;box-shadow:inset 0 0 0 1px rgba(22,74,127,0.12);} 
.canvas-grid,.canvas-notes{display:grid;grid-template-columns:repeat(auto-fit,minmax(220px,1fr));gap:18px;} 
.canvas-card{background:#fff;border-radius:16px;padding:18px;border:1px solid rgba(22,74,127,0.14);box-shadow:0 18px 32px rgba(22,74,127,0.16);} 
.canvas-card h3{margin:0 0 8px;font-size:0.95rem;text-transform:uppercase;letter-spacing:0.08em;color:var(--primary);} 
.canvas-card p{margin:0 0 14px;font-size:0.9rem;color:rgba(31,45,61,0.72);} 
.blank-line{height:26px;border-bottom:1px dashed rgba(22,74,127,0.3);margin-bottom:10px;} 
table.planner{width:100%;border-collapse:collapse;margin-top:18px;background:#fff;border-radius:18px;overflow:hidden;box-shadow:0 22px 38px rgba(22,74,127,0.2);} 
table.planner th,table.planner td{padding:16px;border-bottom:1px solid rgba(22,74,127,0.12);font-size:0.95rem;vertical-align:top;} 
table.planner th{background:rgba(22,74,127,0.15);letter-spacing:0.08em;text-transform:uppercase;color:var(--primary);font-size:0.85rem;} 
.write-cell{min-height:90px;border:1px dashed rgba(22,74,127,0.4);border-radius:12px;background:rgba(244,248,255,0.65);} 
.write-cell.large{min-height:120px;} 
.checklist-cards{display:grid;grid-template-columns:repeat(auto-fit,minmax(220px,1fr));gap:18px;margin-top:24px;} 
.checklist-card{background:#fff;border-radius:16px;padding:18px;border:1px solid rgba(22,74,127,0.14);box-shadow:0 18px 30px rgba(22,74,127,0.16);} 
.data-quality{display:grid;grid-template-columns:repeat(auto-fit,minmax(240px,1fr));gap:18px;} 
.quality-card{background:#fff;border-radius:16px;padding:18px;border:1px solid rgba(22,74,127,0.14);box-shadow:0 18px 30px rgba(22,74,127,0.16);} 
.quality-card ul{list-style:none;margin:0;padding:0;} 
.reflection-prompts{display:grid;grid-template-columns:repeat(auto-fit,minmax(280px,1fr));gap:18px;} 
.reflection-card{background:#fff;border-radius:16px;padding:18px;border:1px solid rgba(22,74,127,0.14);box-shadow:0 18px 30px rgba(22,74,127,0.16);} 
.reflection-card textarea{width:100%;border:1px dashed rgba(22,74,127,0.35);border-radius:12px;min-height:140px;background:rgba(244,248,255,0.7);padding:12px;font-family:inherit;font-size:1rem;} 
.sprint-checklist{margin-top:24px;padding:18px;border-radius:18px;background:rgba(22,74,127,0.1);box-shadow:inset 0 0 0 1px rgba(22,74,127,0.16);} 
.sprint-checklist ul{list-style:none;margin:0;padding:0;display:grid;gap:10px;} 
.sprint-checklist label{display:flex;align-items:center;gap:12px;font-size:1rem;} 
.sprint-checklist input{width:18px;height:18px;accent-color:#164a7f;} 
.sketch-space{margin-top:30px;height:340px;border:2px dashed rgba(22,74,127,0.35);border-radius:18px;background:repeating-linear-gradient(90deg,rgba(22,74,127,0.05),rgba(22,74,127,0.05) 20px,transparent 20px,transparent 40px);} 
footer{text-align:center;margin:40px 0 20px;font-size:0.9rem;color:rgba(31,45,61,0.6);} 
@media(max-width:720px){.content,.step-page{padding:32px 22px;} .step-body{flex-direction:column;}}
.theory-canvas{margin:20px 0;}.toc-flow{display:flex;flex-direction:column;gap:20px;margin-bottom:20px;}.toc-row{display:flex;justify-content:center;align-items:center;gap:5px;}.toc-box{background:linear-gradient(135deg,#e3f2fd,#f3e5f5);border:2px solid #2196f3;border-radius:12px;padding:20px;min-width:120px;text-align:center;box-shadow:0 4px 8px rgba(0,0,0,0.1);flex:1;}.toc-title{font-weight:bold;color:#0d47a1;margin-bottom:5px;}.toc-desc{font-size:0.9em;color:#424242;margin-bottom:10px;}.toc-lines .toc-line{height:24px;border-bottom:1px solid #ddd;margin:5px 0;}.toc-arrow{font-size:18px;color:#ff9800;font-weight:bold;}.toc-accent{text-align:center;margin-top:20px;}.toc-accent-dots span{display:inline-block;width:6px;height:6px;background:#4caf50;border-radius:50%;margin:0 2px;}@media print{.toc-box{box-shadow:none;border-width:1px;}}
@media print{@page{margin:15mm;}@page:first{margin:0;} body{background:#fff;} .step-page{page-break-after:always;}}
"""

FACILITATOR_HTML_CSS = """
:root{--primary:#143a63;--accent:#f6aa50;--ink:#1f2d3d;}
*{box-sizing:border-box;}
body{margin:0;font-family:'Segoe UI','Helvetica Neue',Arial,sans-serif;background:linear-gradient(150deg,#f7faff,#ffffff);color:var(--ink);} 
main{margin:0;padding:0;background:#ffffff;} 
.cover-page{width:100vw;height:100vh;display:flex;align-items:center;justify-content:center;background:#ffffff;page-break-after:always;} 
cover-page img{width:100%;height:100%;object-fit:cover;} 
.content{max-width:1024px;margin:0 auto;padding:48px 54px;} 
h1{margin:0 0 14px;color:var(--primary);font-size:2.25rem;letter-spacing:0.05em;} 
h2{margin:34px 0 12px;color:var(--primary);font-size:0.95rem;letter-spacing:0.08em;text-transform:uppercase;} 
p{line-height:1.6;margin:0 0 1em;} 
.card{margin:24px 0;padding:22px 26px;border-radius:20px;background:rgba(20,58,99,0.06);box-shadow:inset 0 0 0 1px rgba(20,58,99,0.12);} 
.table-wrap{margin:26px 0;border-radius:20px;overflow:hidden;box-shadow:0 20px 36px rgba(20,58,99,0.18);} 
table{width:100%;border-collapse:collapse;background:#fff;} 
th,td{padding:16px;border-bottom:1px solid rgba(20,58,99,0.12);font-size:0.92rem;text-align:left;} 
th{background:rgba(20,58,99,0.14);letter-spacing:0.06em;text-transform:uppercase;color:var(--primary);font-size:0.82rem;} 
.flow-section{margin:30px 0;padding:26px 28px;border-radius:22px;background:linear-gradient(135deg,rgba(20,58,99,0.08),rgba(246,170,80,0.16));box-shadow:inset 0 0 0 1px rgba(20,58,99,0.14);} 
.flow-section ul{margin:10px 0 0;padding-left:1.2rem;} 
.prompt-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(230px,1fr));gap:18px;margin-top:18px;} 
.prompt-card{background:#fff;border-radius:16px;padding:18px;border:1px solid rgba(20,58,99,0.12);box-shadow:0 18px 32px rgba(20,58,99,0.16);} 
.prompt-card strong{display:block;margin-bottom:8px;color:var(--primary);font-size:0.85rem;text-transform:uppercase;letter-spacing:0.08em;} 
footer{text-align:center;margin:40px 0 24px;font-size:0.88rem;color:rgba(31,45,61,0.6);} 
@media(max-width:720px){.content{padding:32px 22px;} .flow-section{padding:24px;}}
@media print{@page{margin:14mm;}@page:first{margin:0;} body{background:#fff;} .flow-section,.card{box-shadow:none;}}
"""

def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)

def configure_cover_section(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    zero_margin = Pt(0)
    section.top_margin = section.bottom_margin = zero_margin
    section.left_margin = section.right_margin = zero_margin

def add_content_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

def apply_table_style(table) -> None:
    for style in (PREFERRED_TABLE_STYLE, FALLBACK_TABLE_STYLE):
        try:
            table.style = style
            break
        except KeyError:
            continue

def shade_cell(cell, color: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:val"), "clear")

def bold_center(cell) -> None:
    para = cell.paragraphs[0]
    if para.text:
        para.runs[0].bold = True
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_spacing(paragraph, before: int = 0, after: int = 6) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)

def add_bullet(container, text: str, indent_inches: float = 0.28):
    para = container.add_paragraph(f"- {text}")
    para.paragraph_format.left_indent = Inches(indent_inches)
    add_spacing(para, after=4)
    return para

def add_step_section(doc: Document, step: dict) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    apply_table_style(table)
    header_cell = table.cell(0, 0)
    header_cell.merge(table.cell(0, 1))
    heading = header_cell.paragraphs[0]
    heading.text = f"STEP {step['number']}: {step['title']}"
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(24)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_spacing(heading, after=6)
    shade_cell(header_cell, "D9E8FF")
    instructions = table.cell(1, 0)
    instructions.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    intro = instructions.paragraphs[0]
    intro.text = step["goal"]
    intro.runs[0].bold = True
    intro.runs[0].font.size = Pt(13)
    add_spacing(intro, after=6)
    for action in step["actions"]:
        bullet = add_bullet(instructions, action)
        bullet.runs[0].font.size = Pt(12)
    tip_text = step.get("tip")
    if tip_text:
        tip = instructions.add_paragraph(f"Tip: {tip_text}")
        tip.runs[0].italic = True
        tip.runs[0].font.size = Pt(11)
        add_spacing(tip, before=6, after=0)
    notes_cell = table.cell(1, 1)
    notes_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    shade_cell(notes_cell, "F2F6FF")
    title = notes_cell.paragraphs[0]
    title.text = "Your notes"
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(12)
    add_spacing(title, after=6)
    for index, note in enumerate(step["notes"]):
        label = note.get("label")
        lines = note.get("lines", 2)
        if label:
            label_para = notes_cell.add_paragraph(f"{label}:")
            label_para.runs[0].bold = True
            label_para.runs[0].font.size = Pt(11)
            add_spacing(label_para, after=2)
        for _ in range(lines):
            line = notes_cell.add_paragraph("____________________________")
            add_spacing(line, after=2)
        if index != len(step["notes"]) - 1:
            add_spacing(notes_cell.add_paragraph(""), after=0)

def add_theory_of_change_canvas(doc: Document) -> None:
    doc.add_paragraph("Use this canvas to sketch the logic. Add arrows, highlight assumptions, and keep phrases short.")
    grid = doc.add_table(rows=2, cols=3)
    apply_table_style(grid)
    for cell, data in zip(grid._cells, TOC_GRID_CELLS):
        heading = cell.paragraphs[0]
        heading.text = data["title"]
        heading.runs[0].bold = True
        add_spacing(heading, after=4)
        prompt = cell.add_paragraph(data["prompt"])
        prompt.runs[0].italic = True
        add_spacing(prompt, after=6)
        for _ in range(data["lines"]):
            line = cell.add_paragraph("_______________________________")
            add_spacing(line, after=3)
    notes = doc.add_table(rows=1, cols=2)
    apply_table_style(notes)
    for cell, data in zip(notes._cells, TOC_NOTE_CELLS):
        heading = cell.paragraphs[0]
        heading.text = data["title"]
        heading.runs[0].bold = True
        add_spacing(heading, after=4)
        prompt = cell.add_paragraph(data["prompt"])
        prompt.runs[0].italic = True
        add_spacing(prompt, after=6)
        for _ in range(data["lines"]):
            line = cell.add_paragraph("_______________________________")
            add_spacing(line, after=3)

def add_measurement_planner(doc: Document) -> None:
    doc.add_paragraph("Translate your theory into evidence. Define indicators, instruments, timing, and notes.")
    table = doc.add_table(rows=len(MEASUREMENT_CONSTRUCTS) + 1, cols=len(MEASUREMENT_HEADERS))
    apply_table_style(table)
    for idx, header in enumerate(MEASUREMENT_HEADERS):
        cell = table.cell(0, idx)
        cell.text = header
        bold_center(cell)
        shade_cell(cell, "D9E8FF")
    for row_idx, construct in enumerate(MEASUREMENT_CONSTRUCTS, start=1):
        label_cell = table.cell(row_idx, 0)
        label_cell.text = construct["label"]
        label_cell.paragraphs[0].runs[0].bold = True
        for col in range(1, len(MEASUREMENT_HEADERS)):
            table.cell(row_idx, col).text = "\n" * construct.get("rows", 3)
    randomization = doc.add_table(rows=2, cols=len(RANDOMIZATION_HEADERS))
    apply_table_style(randomization)
    for idx, header in enumerate(RANDOMIZATION_HEADERS):
        cell = randomization.cell(0, idx)
        cell.text = header
        bold_center(cell)
        shade_cell(cell, "D9E8FF")
    for idx in range(len(RANDOMIZATION_HEADERS)):
        randomization.cell(1, idx).text = "\n\n\n"
    checklist = doc.add_table(rows=1, cols=len(CHECKLIST_BLOCKS))
    apply_table_style(checklist)
    for cell, block in zip(checklist._cells, CHECKLIST_BLOCKS):
        title = cell.paragraphs[0]
        title.text = block[0]
        title.runs[0].bold = True
        add_spacing(title, after=4)
        prompt = cell.add_paragraph(block[1])
        prompt.runs[0].italic = True
        add_spacing(prompt, after=6)
        for _ in range(3):
            line = cell.add_paragraph("_______________________________")
            add_spacing(line, after=3)

def add_data_quality_panel(doc: Document) -> None:
    doc.add_paragraph("Use this panel to keep implementation practical and visible.")
    table = doc.add_table(rows=len(DATA_QUALITY_ROWS), cols=2)
    apply_table_style(table)
    for row_idx, row in enumerate(DATA_QUALITY_ROWS):
        title_cell = table.cell(row_idx, 0)
        title_cell.text = row["title"]
        title_cell.paragraphs[0].runs[0].bold = True
        items_cell = table.cell(row_idx, 1)
        items_cell.paragraphs[0].text = ""
        for item in row["items"]:
            add_bullet(items_cell, f"[ ] {item}", indent_inches=0.18)

def add_reflection_space(doc: Document) -> None:
    doc.add_heading("Sprint Reflection", level=2)
    doc.add_paragraph("Capture final insights before you share back with the group.")
    for prompt in REFLECTION_PROMPTS:
        label = doc.add_paragraph(f"{prompt['label']}:")
        label.runs[0].bold = True
        add_spacing(label, after=4)
        for _ in range(prompt.get("lines", 3)):
            line = doc.add_paragraph("____________________________________________________________")
            add_spacing(line, after=3)
    doc.add_paragraph()
    doc.add_heading("Sprint Checklist", level=2)
    for item in SPRINT_CHECKLIST:
        add_bullet(doc, f"[ ] {item}")
    doc.add_paragraph()
    doc.add_heading("Sketch Space", level=2)
    sketch = doc.add_table(rows=1, cols=1)
    apply_table_style(sketch)
    cell = sketch.cell(0, 0)
    for _ in range(14):
        cell.add_paragraph()

def add_session_snapshot(doc: Document) -> None:
    doc.add_heading("Session Snapshot", level=2)
    table = doc.add_table(rows=1, cols=4)
    apply_table_style(table)
    headers = ["Segment", "Focus", "Format", "Time"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        bold_center(cell)
        shade_cell(cell, "D9E8FF")
    for segment, focus, format_, duration in SESSION_SNAPSHOT_ROWS:
        cells = table.add_row().cells
        cells[0].text = segment
        cells[1].text = focus
        cells[2].text = format_
        cells[3].text = duration

def add_facilitator_materials(doc: Document) -> None:
    doc.add_heading("Prep Checklist", level=2)
    for item in FACILITATOR_PREP:
        add_bullet(doc, item)

def add_facilitator_script(doc: Document) -> None:
    doc.add_heading("Facilitator Flow", level=2)
    for title, bullets in FACILITATOR_FLOW:
        heading = doc.add_paragraph(title, style="Heading 3")
        add_spacing(heading, after=4)
        for bullet in bullets:
            add_bullet(doc, bullet)
    doc.add_heading("Coaching Prompts", level=2)
    for prompt in COACHING_PROMPTS:
        add_bullet(doc, prompt)

def add_facilitator_wrap(doc: Document) -> None:
    doc.add_heading("Debrief and Follow-Up", level=2)
    for item in DEBRIEF_BULLETS:
        add_bullet(doc, item)
    doc.add_heading("Quick Reference", level=2)
    table = doc.add_table(rows=1, cols=2)
    apply_table_style(table)
    headers = ["Cue", "Action"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        bold_center(cell)
        shade_cell(cell, "D9E8FF")
    for cue, action in QUICK_REFERENCE_CUES:
        cells = table.add_row().cells
        cells[0].text = cue
        cells[1].text = action

def render_note_blocks(notes: list) -> str:
    blocks = []
    color_palettes = [
        {"bg1": "#e0f7fa", "bg2": "#f7e8ff", "border": "#b6b6d6", "label": "#2a4d69"},
        {"bg1": "#ffe082", "bg2": "#fff9c4", "border": "#ffe082", "label": "#a67c00"},
        {"bg1": "#d1c4e9", "bg2": "#f3e5f5", "border": "#b39ddb", "label": "#4527a0"},
        {"bg1": "#ffcdd2", "bg2": "#ffe0b2", "border": "#ffab91", "label": "#c62828"},
        {"bg1": "#c8e6c9", "bg2": "#e0f2f1", "border": "#81c784", "label": "#388e3c"},
    ]
    for idx, note in enumerate(notes):
        label = html.escape(note.get("label", ""))
        lines = max(1, note.get("lines", 2))
        palette = color_palettes[idx % len(color_palettes)]
        style = (
            f"--note-bg1: {palette['bg1']}; --note-bg2: {palette['bg2']}; "
            f"--note-border: {palette['border']}; --note-label: {palette['label']};"
        )
        if lines <= 2:
            field = '<input type="text" class="note-input" style="height:42px;">'
        else:
            rows = min(8, max(3, lines * 2))
            field = f'<textarea class="note-area" rows="{rows}" style="min-height:{rows*24}px;"></textarea>'
        parts = [f'      <div class="note-card" style="{style}">']
        if label:
            parts.append(f'        <div class="note-label">{label}</div>')
        parts.append(f"        {field}")
        parts.append('      </div>')
        blocks.append('\n'.join(parts))
    return '\n'.join(blocks)

def render_step_html(step: dict, extension_html: str = "") -> str:
    actions_html = '\n'.join(f"          <li>{html.escape(action)}</li>" for action in step["actions"])
    notes_html = render_note_blocks(step["notes"])
    tip_text = step.get("tip")
    tip_html = f'        <div class="tip">Tip: {html.escape(tip_text)}</div>' if tip_text else ""
    extension_block = f"      <div class=\"step-extension\">\n{extension_html}\n      </div>\n" if extension_html else ""
    return (
        f"  <section class=\"step-page\" data-step=\"Step {step['number']}\">\n"
        "    <div class=\"step-wrapper\">\n"
        f"      <div class=\"step-header\">\n        <h2 class=\"step-title\">Step {step['number']}: {html.escape(step['title'])}</h2>\n"
        f"        <div class=\"step-goal\">{html.escape(step['goal'])}</div>\n      </div>\n"
        "      <div class=\"step-body\">\n        <div class=\"step-actions\">\n          <ul>\n"
        + actions_html
        + "\n          </ul>\n"
        + tip_html
        + "\n        </div>\n        <div class=\"note-stack\">\n"
        + notes_html
        + "\n        </div>\n      </div>\n"
        + extension_block
        + "    </div>\n  </section>"
    )

def render_toc_canvas_html() -> str:
    return '''
    <h2 class="section-title">Theory of Change</h2>
    <div class="theory-canvas">
        <div class="toc-flow">
            <div class="toc-row">
                <div class="toc-box needs">
                    <div class="toc-title">1. Needs</div>
                    <div class="toc-desc">The needs related to our stakeholder group.</div>
                    <div class="toc-lines">
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                    </div>
                </div>
                <div class="toc-arrow">→</div>
                <div class="toc-box stakeholder">
                    <div class="toc-title">2. Stakeholder</div>
                    <div class="toc-desc">Who is our key stakeholder group?</div>
                    <div class="toc-lines">
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                    </div>
                </div>
                <div class="toc-arrow">→</div>
                <div class="toc-box activities">
                    <div class="toc-title">3. Activities</div>
                    <div class="toc-desc">Our operational model.</div>
                    <div class="toc-lines">
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                    </div>
                </div>
            </div>
            <div class="toc-row">
                <div class="toc-box preconditions">
                    <div class="toc-title">4. Preconditions</div>
                    <div class="toc-desc">What needs to be in place to ensure our success?</div>
                    <div class="toc-lines">
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                    </div>
                </div>
                <div class="toc-arrow">→</div>
                <div class="toc-box intermediate">
                    <div class="toc-title">5. Intermediate Outcomes</div>
                    <div class="toc-desc small-text">(Short-term behavioral changes that result from the outputs)</div>
                    <div class="toc-lines">
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                    </div>
                </div>
                <div class="toc-arrow">→</div>
                <div class="toc-box impact">
                    <div class="toc-title">6. Impact</div>
                    <div class="toc-desc">Longer term positive changes that we intend to achieve.</div>
                    <div class="toc-lines">
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                        <div class="toc-line"></div>
                    </div>
                </div>
            </div>
        </div>
        <div class="toc-accent">
            <div class="toc-accent-dots">
                <span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span>
            </div>
        </div>
    </div>
    '''

def render_measurement_html() -> str:
    header_html = '\n'.join(f"            <th>{html.escape(h)}</th>" for h in MEASUREMENT_HEADERS)
    body_rows = []
    for construct in MEASUREMENT_CONSTRUCTS:
        cells = [f"            <td><strong>{html.escape(construct['label'])}</strong></td>"]
        for _ in range(len(MEASUREMENT_HEADERS) - 1):
            size = 'write-cell large' if construct.get('rows', 3) >= 4 else 'write-cell'
            cells.append(f"            <td><div class=\"{size}\"></div></td>")
        body_rows.append("          <tr>\n" + '\n'.join(cells) + "\n          </tr>")
    rand_headers = '\n'.join(f"            <th>{html.escape(h)}</th>" for h in RANDOMIZATION_HEADERS)
    rand_cells = '\n'.join('            <td><div class="write-cell large"></div></td>' for _ in RANDOMIZATION_HEADERS)
    checklist_cards = '\n'.join(
        f"        <div class=\"checklist-card\">\n          <h3>{html.escape(title)}</h3>\n          <p>{html.escape(desc)}</p>\n          <div class=\"blank-line\"></div>\n          <div class=\"blank-line\"></div>\n          <div class=\"blank-line\"></div>\n        </div>"
        for title, desc in CHECKLIST_BLOCKS
    )
    return (
        "        <h3>Measurement and Randomization Planner</h3>\n        <table class=\"planner\">\n          <thead>\n            <tr>\n"
        + header_html
        + "\n            </tr>\n          </thead>\n          <tbody>\n"
        + '\n'.join(body_rows)
        + "\n          </tbody>\n        </table>\n        <table class=\"planner\" style=\"margin-top:26px;\">\n          <thead>\n            <tr>\n"
        + rand_headers
        + "\n            </tr>\n          </thead>\n          <tbody>\n            <tr>\n"
        + rand_cells
        + "\n            </tr>\n          </tbody>\n        </table>\n        <div class=\"checklist-cards\">\n"
        + checklist_cards
        + "\n        </div>"
    )

def render_data_quality_html() -> str:
    cards = []
    for row in DATA_QUALITY_ROWS:
        items = '\n'.join(
            f"          <li><label><input type=\"checkbox\"> {html.escape(item)}</label></li>"
            for item in row['items']
        )
        cards.append(
            f"        <div class=\"quality-card\">\n          <h4>{html.escape(row['title'])}</h4>\n          <ul>\n{items}\n          </ul>\n        </div>"
        )
    return (
        "        <h3>Implementation Rhythm</h3>\n        <div class=\"data-quality\">\n"
        + '\n'.join(cards)
        + "\n        </div>"
    )

def render_reflection_html() -> str:
    cards = []
    for prompt in REFLECTION_PROMPTS:
        cards.append(
            f"        <div class=\"reflection-card\">\n          <h3>{html.escape(prompt['label'])}</h3>\n          <textarea rows=\"{max(4, prompt.get('lines', 3) * 3)}\"></textarea>\n        </div>"
        )
    checklist_items = '\n'.join(
        f"          <li><label><input type=\"checkbox\"> {html.escape(item)}</label></li>"
        for item in SPRINT_CHECKLIST
    )
    return (
        "        <h3>Capture and Commit</h3>\n        <div class=\"reflection-prompts\">\n"
        + '\n'.join(cards)
        + "\n        </div>\n        <div class=\"sprint-checklist\">\n          <h3>Final Checklist</h3>\n          <ul>\n"
        + checklist_items
        + "\n          </ul>\n        </div>\n        <div class=\"sketch-space\" title=\"Sketch space for diagrams\"></div>"
    )

def render_session_snapshot_table() -> str:
    header = "      <tr>" + ''.join(f"<th>{html.escape(h)}</th>" for h in ["Segment", "Focus", "Format", "Time"]) + "</tr>"
    rows = '\n'.join(
        "      <tr>" + ''.join(f"<td>{html.escape(cell)}</td>" for cell in row) + "</tr>"
        for row in SESSION_SNAPSHOT_ROWS
    )
    return (
        "    <div class=\"table-wrap\">\n      <table>\n        <thead>\n"
        + header
        + "\n        </thead>\n        <tbody>\n"
        + rows
        + "\n        </tbody>\n      </table>\n    </div>"
    )

def render_quick_reference_table() -> str:
    header = "      <tr><th>Cue</th><th>Action</th></tr>"
    rows = '\n'.join(
        f"      <tr><td>{html.escape(cue)}</td><td>{html.escape(action)}</td></tr>"
        for cue, action in QUICK_REFERENCE_CUES
    )
    return (
        "    <div class=\"table-wrap\">\n      <table>\n        <thead>\n"
        + header
        + "\n        </thead>\n        <tbody>\n"
        + rows
        + "\n        </tbody>\n      </table>\n    </div>"
    )


def build_participant_html(base_dir: Path, cover_path: Path) -> Path:
    html_parts = [
        "<!DOCTYPE html>",
        "<html lang=\"en\">",
        "<head>",
        "  <meta charset=\"utf-8\">",
        "  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">",
        "  <title>Design an RCT: Participant Workbook</title>",
        "  <style>",
        PARTICIPANT_HTML_CSS,
        "  </style>",
        "</head>",
        "<body>",
        "<main>",
    ]
    # 1. Title + How to Use (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append('<h1 class="section-title">Design an RCT: Participant Workbook</h1>')
    html_parts.append('<p>Work through each step to turn your program concept into a randomized control trial design. Capture notes directly in the spaces provided so you leave the sprint ready to refine your design.</p>')
    html_parts.append('<h2 class="section-subtitle">How to Use This Workbook</h2>')
    html_parts.append('<ul class="guidance">')
    for item in PARTICIPANT_GUIDANCE:
        html_parts.append(f'<li>{html.escape(item)}</li>')
    html_parts.append('</ul></section>')

    # 2. Step 1 (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append(render_step_html(PARTICIPANT_STEPS[0]))
    html_parts.append('</section>')

    # 3. Step 2 (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append(render_step_html(PARTICIPANT_STEPS[1]))
    html_parts.append('</section>')

    # 4. Theory of Change (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append(render_toc_canvas_html())
    html_parts.append('</section>')

    # 5. Step 3 up to instruments (single page)
    html_parts.append('<section class="section-page">')
    step3 = PARTICIPANT_STEPS[2].copy()
    step3_notes = [step3['notes'][0], step3['notes'][1]]
    step3_partial = step3.copy()
    step3_partial['notes'] = step3_notes
    html_parts.append(render_step_html(step3_partial))
    html_parts.append('</section>')

    # 6. Step 3 (measurement to power calculations) (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append(render_measurement_html())
    html_parts.append('</section>')

    # 7. Data quality to Decision trigger (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append(render_data_quality_html())
    html_parts.append('</section>')

    # 8. Step 4 (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append(render_step_html(PARTICIPANT_STEPS[3]))
    html_parts.append('</section>')

    # 9. Step 5 up to Risks to watch (single page)
    html_parts.append('<section class="section-page">')
    step5 = PARTICIPANT_STEPS[4].copy()
    step5_notes = [step5['notes'][0]]
    step5_partial = step5.copy()
    step5_partial['notes'] = step5_notes
    html_parts.append(render_step_html(step5_partial))
    html_parts.append('</section>')

    # 10. Implementation rhythm up to risks (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append(render_data_quality_html())
    html_parts.append('</section>')

    # 11. Step 6 up to Next steps (single page)
    html_parts.append('<section class="section-page">')
    step6 = PARTICIPANT_STEPS[5].copy()
    step6_notes = [step6['notes'][0], step6['notes'][1], step6['notes'][2]]
    step6_partial = step6.copy()
    step6_partial['notes'] = step6_notes
    html_parts.append(render_step_html(step6_partial))
    html_parts.append('</section>')

    # 12. Capture and commit (single page)
    html_parts.append('<section class="section-page">')
    html_parts.append(render_reflection_html())
    html_parts.append('</section>')

    html_parts.extend([
        '<footer>Bringing your program concept into an RCT Design!</footer>',
        '</main>',
        '</body>',
        '</html>',
    ])
    output_path = base_dir / "RCT_Design_Participant_Booklet.html"
    output_path.write_text('\n'.join(html_parts), encoding="utf-8")
    return output_path

def build_facilitator_html(base_dir: Path, cover_path: Path) -> Path:
    cover_name = cover_path.name
    html_parts = [
        "<!DOCTYPE html>",
        "<html lang=\"en\">",
        "<head>",
        "  <meta charset=\"utf-8\">",
        "  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">",
        "  <title>Design an RCT: Facilitator Guide</title>",
        "  <style>",
        FACILITATOR_HTML_CSS,
        "  </style>",
        "</head>",
        "<body>",
        "<main>",
        f"  <section class=\"cover-page\"><img src=\"{html.escape(cover_name)}\" alt=\"Design an RCT cover\"></section>",
        "  <div class=\"content\">",
        "    <h1>Design an RCT: Facilitator Guide</h1>",
        "    <p>Use this guide as a quick-reference companion to the participant workbook. Keep the cadence tight, surface assumptions, and encourage teams to document decisions clearly.</p>",
        render_session_snapshot_table(),
        "    <div class=\"card\">",
        "      <h2>Prep Checklist</h2>",
        "      <ul>",
    ]
    html_parts.extend(f"        <li>{html.escape(item)}</li>" for item in FACILITATOR_PREP)
    html_parts.append("      </ul>\n    </div>")
    html_parts.append("    <section class=\"flow-section\">\n      <h2>Facilitator Flow</h2>")
    for title, bullets in FACILITATOR_FLOW:
        html_parts.append(f"      <h3>{html.escape(title)}</h3>\n      <ul>")
        html_parts.extend(f"        <li>{html.escape(bullet)}</li>" for bullet in bullets)
        html_parts.append("      </ul>")
    html_parts.append("    </section>")
    html_parts.append("    <div class=\"card\">\n      <h2>Coaching Prompts</h2>\n      <ul>")
    html_parts.extend(f"        <li>{html.escape(prompt)}</li>" for prompt in COACHING_PROMPTS)
    html_parts.append("      </ul>\n    </div>")
    html_parts.append("    <div class=\"card\">\n      <h2>Debrief and Follow-Up</h2>\n      <ul>")
    html_parts.extend(f"        <li>{html.escape(item)}</li>" for item in DEBRIEF_BULLETS)
    html_parts.append("      </ul>\n    </div>")
    html_parts.append(render_quick_reference_table())
    html_parts.extend([
        "    <footer>Print or save alongside the workbook to keep your facilitation tight and repeatable.</footer>",
        "  </div>",
        "</main>",
        "</body>",
        "</html>",
    ])
    output_path = base_dir / "RCT_Design_Facilitator_Guide.html"
    output_path.write_text('\n'.join(html_parts), encoding="utf-8")
    return output_path

# --- Final function block: CLI entry point ---

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Generate RCT participant booklet and facilitator guide HTML files.")
    parser.add_argument("--output-dir", type=str, default=".", help="Directory to save output files")
    parser.add_argument("--cover", type=str, required=True, help="Path to cover image file")
    parser.add_argument("--pdf", action="store_true", help="Also generate PDF output")
    parser.add_argument("--docx", action="store_true", help="Also generate DOCX output")
    args = parser.parse_args()

    base_dir = Path(args.output_dir)
    cover_path = Path(args.cover)
    base_dir.mkdir(parents=True, exist_ok=True)

    participant_html = build_participant_html(base_dir, cover_path)
    facilitator_html = build_facilitator_html(base_dir, cover_path)

    print(f"Participant booklet saved to: {participant_html}")
    print(f"Facilitator guide saved to: {facilitator_html}")

    # PDF output
    if args.pdf:
        try:
            import pdfkit
            pdf_path = participant_html.with_suffix('.pdf')
            pdfkit.from_file(str(participant_html), str(pdf_path))
            print(f"Participant booklet PDF saved to: {pdf_path}")
        except Exception as e:
            print(f"PDF generation failed: {e}\nInstall pdfkit and wkhtmltopdf for PDF support.")

    # DOCX output
    if args.docx:
        try:
            from bs4 import BeautifulSoup, Tag
            docx_path = participant_html.with_suffix('.docx')
            soup = BeautifulSoup(participant_html.read_text(encoding="utf-8"), "html.parser")
            doc = Document()
            for section in soup.find_all('section', class_='section-page'):
                for el in list(section.contents):
                    if not isinstance(el, Tag):
                        continue
                    if el.name == 'h1':
                        doc.add_heading(el.get_text(), level=1)
                    elif el.name == 'h2':
                        doc.add_heading(el.get_text(), level=2)
                    elif el.name == 'h3':
                        doc.add_heading(el.get_text(), level=3)
                    elif el.name == 'h4':
                        doc.add_heading(el.get_text(), level=4)
                    elif el.name == 'p':
                        doc.add_paragraph(el.get_text())
                    elif el.name == 'ul':
                        for li in el.find_all('li', recursive=False):
                            doc.add_paragraph(li.get_text(), style='List Bullet')
                doc.add_page_break()
            doc.save(str(docx_path))
            print(f"Participant booklet DOCX saved to: {docx_path}")
        except Exception as e:
            print(f"DOCX generation failed: {e}\nInstall python-docx and beautifulsoup4 for DOCX support.")


if __name__ == "__main__":
    main()

